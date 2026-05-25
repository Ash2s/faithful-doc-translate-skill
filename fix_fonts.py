# -*- coding: utf-8 -*-
"""
Fix fonts and spacing on existing EN .docx files.
Post-processing utility for the immigration-doc-translate skill.

Walks all .docx files under BASE_EN and forces:
- Times New Roman 12pt on all runs (preserving bold/highlight/underline)
- 2x line spacing for body paragraphs, 1x for table cells
"""
import sys, os
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from translate_engine import save_doc


def fix_run_font(run):
    """Force a single run to Times New Roman 12pt, preserving bold/italic/highlight."""
    rPr = run._r.find(qn('w:rPr'))
    if rPr is None:
        rPr = OxmlElement('w:rPr')
        run._r.insert(0, rPr)

    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:ascii'), 'Times New Roman')
    rFonts.set(qn('w:hAnsi'), 'Times New Roman')
    rFonts.set(qn('w:eastAsia'), 'Times New Roman')

    sz = rPr.find(qn('w:sz'))
    if sz is None:
        sz = OxmlElement('w:sz')
        rPr.append(sz)
    sz.set(qn('w:val'), '24')  # 12pt

    szCs = rPr.find(qn('w:szCs'))
    if szCs is None:
        szCs = OxmlElement('w:szCs')
        rPr.append(szCs)
    szCs.set(qn('w:val'), '24')

    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)


def set_para_spacing(para, double_space):
    """Set paragraph line spacing: 2x or 1x."""
    pPr = para._p.get_or_add_pPr()
    sp = pPr.find(qn('w:spacing'))
    if sp is None:
        sp = OxmlElement('w:spacing')
        pPr.append(sp)
    sp.set(qn('w:line'), '480' if double_space else '240')
    sp.set(qn('w:lineRule'), 'auto')


def is_heading(para):
    """Check if a paragraph uses a Heading style."""
    style = para.style
    if style is None:
        return False
    name = style.name or ''
    return name.startswith('Heading')


def fix_doc_fonts(path):
    """Fix all runs in a document: TNR 12pt, with spacing."""
    doc = Document(path)
    modified = False

    for para in doc.paragraphs:
        if not para.runs:
            continue
        is_hdg = is_heading(para)
        for run in para.runs:
            fix_run_font(run)
            modified = True
        if not is_hdg:
            set_para_spacing(para, True)

    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        fix_run_font(run)
                        modified = True
                    set_para_spacing(para, False)

    if modified:
        save_doc(doc, path)
    return modified
