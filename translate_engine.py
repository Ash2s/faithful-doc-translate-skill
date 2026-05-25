# -*- coding: utf-8 -*-
"""
EB1A Translation Engine v4 (Per-Run Formatting)
- Longest-match term replacement (avoids partial consumption)
- Version-number fuzzy matching: "系统V1.0" matches "系统 V1.0"
- Auto-translate common proper nouns not in glossary
- Copy source docx XML, replace text in-place → preserves highlight/bold/underline/alignment
- Strips background shading (w:shd) — white paper, black text only
- Forces Times New Roman 12pt at all levels: theme XML (Hans script), style defaults, run fonts
- Strips theme font references (asciiTheme/hAnsiTheme/cstheme) → explicit font names
- Chinese residual verification
- Paragraph-level CN bracket-label replacement (handles split-run 【】tags)
- Per-run replacement: each run keeps its own formatting (not consensus collapse)

Design principle: This engine handles FORMATTING only.  Translation is done
by Claude in-conversation.  The engine provides glossary matching for known
terms, but full-sentence AI translation is NOT embedded here — it belongs
in the Claude conversation so any model/provider can be used.
"""
import sys, os, re, shutil, zipfile
from typing import Optional
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import lxml.etree as etree


# ============================================================
# Common proper nouns: auto-translated when encountered
# ============================================================
COMMON_TERMS = {
    # Companies / orgs
    '中海地产': 'China Overseas Land & Investment (COOP)',
    '万科地产': 'Vanke Real Estate',
    '景瑞控股': 'Jingrui Holdings',
    '华润': 'China Resources',
    '百安居': 'B&Q',
    '旭辉': 'Xuhui Group',
    '招商': 'China Merchants',
    '上海地产': 'Shanghai Land',
    '首旅如家': 'BTU Ruijia',
    '锦江': 'Jinjiang Hotels',
    '华住': 'Huazhu Hotels',
    '格林': 'GreenTree Hotels',
    '恒大': 'Evergrande',
    '万科': 'Vanke',
    '中海': 'China Overseas Land & Investment (COOP)',
    '清华': 'Tsinghua University',
    '长安大学': "Chang'an University",
    '中欧国际商学院': 'China Europe International Business School (CEIBS)',
    '香港技术研究院': 'Hong Kong Institute of Technology',
    '中国智慧工程研究会': 'China Society of Intelligent Engineering',

    # Awards / competitions
    '创意空间杯': 'Creative Space Cup',

    # Technical terms
    '建筑声学': 'architectural acoustics',
    '数字化转型': 'digital transformation',
    '双碳': '"dual carbon"',
    '碳达峰': 'carbon peak',
    '碳中和': 'carbon neutrality',
    '智慧城市': 'smart city',
    '零碳建筑': 'zero-carbon buildings',
    '装配式内装': 'prefabricated interior decoration',
    '工业化': 'industrialization',
    '内装工业化': 'interior industrialization',
    '数字化': 'digitalization',
    '数字化工具': 'digital tools',
    '生物气候建筑': 'bioclimatic architecture',
    '室内热舒适性': 'indoor thermal comfort',
    '拓扑算法': 'topological algorithms',
    '空间句法': 'spatial syntax',
    '绿色决策支持系统': 'green decision-support system',
    '系统架构师': 'technical architect',
    '绿色建材': 'green building materials',
    '装修': 'decoration',
    '精装': 'finished interior',
    '装配式建筑': 'prefabricated building',
    '全装修': 'full-decoration',
    'SI体系': 'SI structural system',
    '现场湿作业': 'on-site wet work',
    '无醛无毒': 'formaldehyde-free and non-toxic',
    '资源回收再利用': 'recycling and reuse',
    '装修垃圾': 'decoration waste',
    '部品': 'components',
    '首席执行官': 'Chief Executive Officer',
    '云平台': 'Cloud Platform',

    # Titles / roles
    '注册建筑师': 'Registered Architect',
    '国家高新技术企业': 'National High-Tech Enterprise',
    '参编': 'participated in drafting',
    '主编': 'editor-in-chief',
    '第一负责人': 'Principal Investigator',
    '联合创始人': 'Co-Founder',
    '核心合伙人': 'Core Partner',
    '客座教授': 'Visiting Professor',
    '特聘专家': 'Distinguished Expert',
    '评审专家': 'Review Expert',
    '副研究员': 'Associate Researcher',

    # Names (personal)
    '向宠': 'Xiang Chong',
    '谭旭阳': 'Tan Xuyang',
    '乐云': 'Le Yun',

    # Brand names from docs
    '卡瑞': 'CARR',
    'Pincloud': 'Pincloud',
    '桔装无忧屋佳': 'Juzhuang Wuyou Wujia',
    '品宅·集栋装饰工程': 'Pinzhai · Jidong Decoration Engineering',
    '品宅·一站美办': 'Pinzhai · Yizhan Meiban',
    '品宅·酒店快益装': 'Pinzhai · Jiuyi Kuaiyi',
    '品宅装配式内装设计研究院': 'Pinzhai Prefabricated Interior Design Research Institute',
}


def normalize(text):
    """Remove spaces for fuzzy matching."""
    return text.replace(' ', '').replace('　', '').replace(' ', '')


def normalize_quotes(s: str) -> str:
    """Convert all straight double quotes to curly quotes for consistent matching.

    Straight quotes (U+0022) in glossary keys or document text cause
    dictionary-lookup failures.  This function converts them to the canonical
    curly form (U+201C/U+201D) that Word natively uses.
    """
    if not s:
        return s
    result = []
    open_quote = True
    for ch in s:
        if ch == '"':
            result.append('“' if open_quote else '”')
            open_quote = not open_quote
        else:
            result.append(ch)
    return ''.join(result)


def load_glossary(xlsx_path: str) -> dict:
    """Load terminology from Excel glossary file."""
    import openpyxl
    wb = openpyxl.load_workbook(xlsx_path)
    ws = wb['术语表'] if '术语表' in wb.sheetnames else wb.active
    glossary = {}
    for row in ws.iter_rows(values_only=True):
        cn, en = row[0], row[1]
        if cn and en and isinstance(cn, str) and isinstance(en, str):
            cn, en = cn.strip(), en.strip()
            if cn.startswith(('zh', 'en', '专利', '公司', '媒体', '机构', '地址', '活动', '其他')):
                continue
            glossary[cn] = en
    return glossary


def build_matcher(glossary: dict) -> list:
    """
    Build sorted (cn, en) list for longest-match replacement.
    Also store normalized version for fuzzy matching.
    Returns: [(cn, en, norm_cn_or_None)]
    """
    candidates = []
    for cn, en in glossary.items():
        cn = normalize_quotes(cn)
        norm = normalize(cn)
        candidates.append((cn, en, norm if norm != cn else None))

    for cn, en in COMMON_TERMS.items():
        cn = normalize_quotes(cn)
        norm = normalize(cn)
        candidates.append((cn, en, norm if norm != cn else None))

    candidates.sort(key=lambda x: -len(x[0]))
    return candidates


def translate_text(text: str, matcher: list) -> tuple:
    """
    Apply longest-match term replacement.
    Returns: (translated_text, list_of_unmatched_chinese_terms)
    """
    if not text:
        return text, []

    text = normalize_quotes(text)
    result = text
    i = 0
    while i < len(result):
        matched = False
        for cn, en, norm_cn in matcher:
            en_clean = en.split('(')[0].strip()

            if result[i:].startswith(cn):
                if result[i:i+len(en)].startswith(en) or (i > 0 and result[i-len(en_clean):i] == en_clean):
                    i += 1
                    continue
                result = result[:i] + en + result[i+len(cn):]
                i += len(en)
                matched = True
                break
            if norm_cn and result[i:].startswith(norm_cn):
                result = result[:i] + en + result[i+len(norm_cn):]
                i += len(en)
                matched = True
                break
        if not matched:
            i += 1

    unmatched = list(dict.fromkeys(re.findall(r'[一-龿]{2,}', result)))
    return result, unmatched


# ============================================================
# Format preservation: copy source XML, replace text in-place
# ============================================================

def _strip_shading(rPr_elem):
    """Remove w:shd child elements from an rPr element (in-place)."""
    if rPr_elem is None:
        return
    for shd in rPr_elem.findall(qn('w:shd')):
        rPr_elem.remove(shd)


def _force_tnr_12pt(rPr_elem):
    """
    Override rPr font to Times New Roman 12pt.
    Keeps other formatting (bold, highlight, underline, color, etc.) intact.
    This is the key difference from just copying source formatting:
    visual style is preserved, but font face and size are always overridden.
    """
    if rPr_elem is None:
        return
    # Override w:rFonts — ascii, hAnsi, eastAsia all to Times New Roman
    rFonts = rPr_elem.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr_elem.insert(0, rFonts)
    rFonts.set(qn('w:ascii'), 'Times New Roman')
    rFonts.set(qn('w:hAnsi'), 'Times New Roman')
    rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    # Override w:sz to 12pt (24 half-points)
    sz = rPr_elem.find(qn('w:sz'))
    if sz is None:
        sz = OxmlElement('w:sz')
        rPr_elem.append(sz)
    sz.set(qn('w:val'), '24')
    # Override w:szCs (complex script size) to 12pt
    szCs = rPr_elem.find(qn('w:szCs'))
    if szCs is None:
        szCs = OxmlElement('w:szCs')
        rPr_elem.append(szCs)
    szCs.set(qn('w:val'), '24')


def replace_runs_in_place(para, matcher: list):
    """
    Replace each run's text using glossary matching, preserving per-run formatting.

    Unlike ``clear_and_set_text`` which collapses all runs into one with consensus
    formatting, this function keeps each run's original formatting (highlight,
    underline, bold, italic, color) intact.  Background shading is removed and
    font is forced to Times New Roman 12pt.

    Args:
        para: python-docx Paragraph object.
        matcher: list of (cn, en, norm) from build_matcher().
    """
    p_elem = para._p

    # Remove paragraph-level shading
    pPr = p_elem.find(qn('w:pPr'))
    if pPr is not None:
        for shd in pPr.findall(qn('w:shd')):
            pPr.remove(shd)

    for r in p_elem.findall(qn('w:r')):
        t_elements = r.findall(qn('w:t'))
        # Skip runs with no text
        has_text = any(t.text and t.text.strip() for t in t_elements)
        if not has_text:
            continue

        # Get original text
        orig_text = ''.join(t.text or '' for t in t_elements)

        # Translate this run's text individually
        trans_text, _ = translate_text(orig_text, matcher)
        if trans_text == orig_text:
            continue  # nothing changed, keep as-is

        # Clear existing <w:t> elements (they will be replaced)
        for t in t_elements:
            r.remove(t)

        # Add new <w:t> with translated text
        new_t = OxmlElement('w:t')
        new_t.set(qn('xml:space'), 'preserve')
        new_t.text = trans_text
        r.append(new_t)

        # Clean up rPr: remove shading, force TNR 12pt, keep everything else
        rPr = r.find(qn('w:rPr'))
        if rPr is None:
            rPr = OxmlElement('w:rPr')
            r.insert(0, rPr)

        _strip_shading(rPr)
        _force_tnr_12pt(rPr)


def clear_and_set_text(para, new_text, force_font=True):
    """
    Clear all runs in a paragraph and set new text.
    Preserves highlight, bold, font formatting using consensus rules:
    - Bold: only if ALL text runs are bold (avoids making entire abstract bold)
    - Highlight: if ANY run has highlight, preserve it (for name highlighting)
    - Italic: only if ALL text runs are italic
    - Underline: if ANY run has underline
    Strips background shading (w:shd) — clean white background only.
    Also strips paragraph-level shading.
    """
    p_elem = para._p
    ns_w = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

    # Remove paragraph-level shading
    pPr = p_elem.find(qn('w:pPr'))
    if pPr is not None:
        for shd in pPr.findall(qn('w:shd')):
            pPr.remove(shd)

    # Collect all runs that contain text
    text_runs = []
    for r in p_elem.findall(qn('w:r')):
        t_elements = r.findall(qn('w:t'))
        has_text = any(t.text and t.text.strip() for t in t_elements if t.text)
        if has_text:
            text_runs.append(r)

    # If no text runs found, use any run with rPr as fallback
    if not text_runs:
        for r in p_elem.findall(qn('w:r')):
            if r.find(qn('w:rPr')) is not None:
                text_runs.append(r)
                break

    # Analyze formatting consensus across all text runs
    all_bold = True
    any_bold = False
    all_italic = True
    any_italic = False
    all_underline = True
    any_underline = False
    any_highlight = False
    highlight_val = None
    best_rPr = None  # The rPr with the most formatting features

    for r in text_runs:
        rPr = r.find(qn('w:rPr'))
        if rPr is None:
            all_bold = False
            all_italic = False
            continue

        # Bold check
        b = rPr.find(qn('w:b'))
        is_bold = b is not None and b.get(qn('w:val')) != '0'
        if is_bold:
            any_bold = True
        else:
            all_bold = False

        # Italic check
        i = rPr.find(qn('w:i'))
        is_italic = i is not None and i.get(qn('w:val')) != '0'
        if is_italic:
            any_italic = True
        else:
            all_italic = False

        # Highlight check
        hl = rPr.find(qn('w:highlight'))
        if hl is not None:
            hl_val = hl.get(qn('w:val'))
            if hl_val and hl_val != 'none':
                any_highlight = True
                highlight_val = hl_val

        # Underline check
        u = rPr.find(qn('w:u'))
        is_underline = u is not None and u.get(qn('w:val')) != 'none'
        if is_underline:
            any_underline = True
        else:
            all_underline = False

        # Track best rPr (most child elements = most formatting)
        if best_rPr is None or len(list(rPr)) > len(list(best_rPr)):
            best_rPr = rPr

    # Remove all existing runs
    for r in p_elem.findall(qn('w:r')):
        p_elem.remove(r)

    # Create new run with appropriate formatting
    new_run = OxmlElement('w:r')
    new_rPr = etree.SubElement(new_run, qn('w:rPr'))

    # Copy base formatting from best rPr, then strip shading and adjust based on consensus
    if best_rPr is not None:
        for child in best_rPr:
            tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
            # Skip shading (always stripped)
            if tag == 'shd':
                continue
            # Skip bold/italic/highlight/underline — we'll set them based on consensus
            if tag in ('b', 'i', 'highlight', 'u'):
                continue
            new_rPr.append(etree.fromstring(etree.tostring(child)))

    # Apply consensus formatting
    if all_bold and any_bold:
        b = etree.SubElement(new_rPr, qn('w:b'))
    if all_italic and any_italic:
        i = etree.SubElement(new_rPr, qn('w:i'))
    if any_highlight and highlight_val:
        hl = etree.SubElement(new_rPr, qn('w:highlight'))
        hl.set(qn('w:val'), highlight_val)
    if any_underline:
        u = etree.SubElement(new_rPr, qn('w:u'))
        u.set(qn('w:val'), 'single')

    # Force Times New Roman 12pt
    if force_font:
        _force_tnr_12pt(new_rPr)
    else:
        # Still ensure font name if not set
        rFonts = new_rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = etree.SubElement(new_rPr, qn('w:rFonts'))
        if not rFonts.get(qn('w:ascii')):
            rFonts.set(qn('w:ascii'), 'Times New Roman')
            rFonts.set(qn('w:hAnsi'), 'Times New Roman')
            rFonts.set(qn('w:eastAsia'), 'Times New Roman')

    t_elem = etree.SubElement(new_run, qn('w:t'))
    t_elem.set(qn('xml:space'), 'preserve')
    t_elem.text = new_text
    p_elem.append(new_run)


def strip_all_shading(doc):
    """Post-process: remove all w:shd from paragraphs and runs in the whole document."""
    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    for shd in doc.element.findall('.//w:shd', ns):
        shd.getparent().remove(shd)


def set_spacing(para, double_space=True):
    """Set paragraph line spacing. True = 2x, False = 1x."""
    pPr = para._p.get_or_add_pPr()
    sp = pPr.find(qn('w:spacing'))
    if sp is None:
        sp = OxmlElement('w:spacing')
        pPr.append(sp)
    sp.set(qn('w:line'), '480' if double_space else '240')
    sp.set(qn('w:lineRule'), 'auto')


# ============================================================
# Font post-processing: eliminate 宋体 at every level
# ============================================================

def _fix_theme_xml(docx_path):
    """Replace 宋体 with Times New Roman in theme font scheme (Hans script)."""
    temp_path = docx_path + '.tmp'
    with zipfile.ZipFile(docx_path, 'r') as zin:
        with zipfile.ZipFile(temp_path, 'w', zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if 'theme' in item.filename.lower():
                    text = data.decode('utf-8')
                    text = text.replace(
                        '<a:font script="Hans" typeface="宋体"/>',
                        '<a:font script="Hans" typeface="Times New Roman"/>'
                    )
                    data = text.encode('utf-8')
                zout.writestr(item, data)
    shutil.move(temp_path, docx_path)


def _fix_run_fonts(doc):
    """Force TNR 12pt on every run in the doc (post-processing safety net)."""
    for para in list(doc.paragraphs) + [p for t in doc.tables for r in t.rows for c in r.cells for p in c.paragraphs]:
        for r in para._p.findall(qn('w:r')):
            rPr = r.find(qn('w:rPr'))
            if rPr is None:
                rPr = etree.SubElement(r, qn('w:rPr'))
            _force_tnr_12pt(rPr)
    # Textboxes
    for txbx in doc.element.findall(f'.//{qn("w:txbxContent")}'):
        for p in txbx.findall(qn('w:p')):
            for r in p.findall(qn('w:r')):
                rPr = r.find(qn('w:rPr'))
                if rPr is None:
                    rPr = etree.SubElement(r, qn('w:rPr'))
                _force_tnr_12pt(rPr)


def _fix_style_fonts(doc):
    """Set docDefaults and Normal style to explicit Times New Roman."""
    styles_elem = doc.styles.element

    # docDefaults
    docDefaults = styles_elem.find(qn('w:docDefaults'))
    if docDefaults is None:
        docDefaults = etree.SubElement(styles_elem, qn('w:docDefaults'))
    rPrDefault = docDefaults.find(qn('w:rPrDefault'))
    if rPrDefault is None:
        rPrDefault = etree.SubElement(docDefaults, qn('w:rPrDefault'))
    rPr = rPrDefault.find(qn('w:rPr'))
    if rPr is None:
        rPr = etree.SubElement(rPrDefault, qn('w:rPr'))
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = etree.SubElement(rPr, qn('w:rFonts'))
    for tag in ('w:ascii', 'w:hAnsi', 'w:eastAsia', 'w:cs'):
        rFonts.set(qn(tag), 'Times New Roman')
    for attr in list(rFonts.attrib):
        if 'theme' in attr.lower():
            del rFonts.attrib[attr]

    # Normal style
    normal = doc.styles['Normal']
    n_rPr = normal.element.find(qn('w:rPr'))
    if n_rPr is None:
        n_rPr = etree.SubElement(normal.element, qn('w:rPr'))
    n_rFonts = n_rPr.find(qn('w:rFonts'))
    if n_rFonts is None:
        n_rFonts = etree.SubElement(n_rPr, qn('w:rFonts'))
    for tag in ('w:ascii', 'w:hAnsi', 'w:eastAsia', 'w:cs'):
        n_rFonts.set(qn(tag), 'Times New Roman')
    for attr in list(n_rFonts.attrib):
        if 'theme' in attr.lower():
            del n_rFonts.attrib[attr]

    # All other styles: clear theme refs, fill in missing fonts
    for style in styles_elem.findall(qn('w:style')):
        style_rPr = style.find(qn('w:rPr'))
        if style_rPr is not None:
            style_rFonts = style_rPr.find(qn('w:rFonts'))
            if style_rFonts is not None:
                for attr in list(style_rFonts.attrib):
                    if 'theme' in attr.lower():
                        del style_rFonts.attrib[attr]
                for tag in ('w:ascii', 'w:hAnsi', 'w:eastAsia', 'w:cs'):
                    style_rFonts.set(qn(tag), 'Times New Roman')


def fix_document_fonts(docx_path):
    """
    Post-process a saved docx to force Times New Roman at every level:
    theme XML → style defaults → run fonts. Eliminates 宋体 from all sources.
    Call this after save_doc().
    """
    _fix_theme_xml(docx_path)
    doc = Document(docx_path)
    _fix_run_fonts(doc)
    _fix_style_fonts(doc)
    doc.save(docx_path)


def save_doc(doc, dst_path):
    """Save docx with long-path workaround on Windows, then fix fonts."""
    abs_dst = os.path.abspath(dst_path)
    save_path = '\\\\?\\' + abs_dst if len(abs_dst) >= 260 else abs_dst
    os.makedirs(os.path.dirname(save_path), exist_ok=True)
    doc.save(save_path)
    fix_document_fonts(save_path)


def copy_and_translate(src_path, dst_path, para_map):
    """
    Copy source docx, translate paragraphs using para_map {index: english_text}.
    Uses clear_and_set_text (consensus formatting, collapses multi-run into one).
    Tables cells can also be mapped via key f'<t>{ti}r{ri}c{ci}'.
    """
    abs_src = os.path.abspath(src_path)
    abs_dst = os.path.abspath(dst_path)
    os.makedirs(os.path.dirname(abs_dst), exist_ok=True)

    doc = Document(src_path)

    # Process paragraphs
    for i, para in enumerate(doc.paragraphs):
        if i in para_map:
            clear_and_set_text(para, para_map[i])
            # Body paragraphs: double spacing
            set_spacing(para, True)

    # Process tables
    for ti, tbl in enumerate(doc.tables):
        for ri, row in enumerate(tbl.rows):
            for ci, cell in enumerate(row.cells):
                key = f't{ti}r{ri}c{ci}'
                if key in para_map:
                    cell_text = para_map[key]
                    for pi, para in enumerate(cell.paragraphs):
                        if pi == 0:
                            clear_and_set_text(para, cell_text)
                            # Table cells: single spacing
                            set_spacing(para, False)
                        else:
                            p_elem = para._p
                            p_elem.getparent().remove(p_elem)

    save_doc(doc, dst_path)


def verify_no_cn(doc_or_path):
    """Check for Chinese residual characters. Returns True if clean."""
    cn_pat = re.compile(r'[一-龿]')
    if isinstance(doc_or_path, str):
        doc = Document(doc_or_path)
    else:
        doc = doc_or_path

    found = False
    for i, para in enumerate(doc.paragraphs):
        if cn_pat.search(para.text):
            txt = para.text.strip()[:100]
            print(f'  CN P{i}: {txt}')
            found = True
    for ti, tbl in enumerate(doc.tables):
        for ri, row in enumerate(tbl.rows):
            for ci, cell in enumerate(row.cells):
                if cn_pat.search(cell.text):
                    txt = cell.text.strip()[:100]
                    print(f'  CN T{ti}R{ri}C{ci}: {txt}')
                    found = True
    # Scan textboxes (not covered by paragraphs/tables)
    for txbx in doc.element.findall(f'.//{qn("w:txbxContent")}'):
        for p in txbx.findall(qn('w:p')):
            p_text = ''.join(t.text or '' for t in p.findall(f'.//{qn("w:t")}'))
            if cn_pat.search(p_text):
                txt = p_text.strip()[:100]
                print(f'  CN TXBX: {txt}')
                found = True
    return not found


# ============================================================
# Post-processing: bracket-label replacement
# ============================================================

CN_LABEL_MAP = {
    '【图片】': '[Image]',
    '【条形码】': '[Barcode]',
    '【图标】': '[Logo]',
    '【盖章】': '[Seal]',
    '【英文材料无需翻译】': '[No Translation Required — Source Material in English]',
}


def fix_cn_labels(doc):
    """
    Paragraph-level replacement of 【】bracket labels → [English].
    Detects at paragraph.text level (not run level) because these labels
    are frequently split across multiple w:r runs in OOXML.
    """
    for para in doc.paragraphs:
        txt = para.text.strip()
        if txt in CN_LABEL_MAP:
            clear_and_set_text(para, CN_LABEL_MAP[txt])
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    txt = para.text.strip()
                    if txt in CN_LABEL_MAP:
                        clear_and_set_text(para, CN_LABEL_MAP[txt])
    # Textboxes
    for txbx in doc.element.findall(f'.//{qn("w:txbxContent")}'):
        for p in txbx.findall(qn('w:p')):
            p_text = ''.join(t.text or '' for t in p.findall(f'.//{qn("w:t")}'))
            if p_text.strip() in CN_LABEL_MAP:
                clear_and_set_text(p, CN_LABEL_MAP[p_text.strip()])


def translate_docx(src_path: str, dst_path: str, glossary_path: str,
                   user_translations: dict = None,
                   uncertain_terms: list = None) -> dict:
    """
    Translate a docx file using glossary + common terms, preserving source format.

    Uses copy-and-replace approach (not create-new-doc) to keep all visual
    formatting (highlight, bold, underline, alignment).

    Args:
        src_path: source CN docx
        dst_path: destination EN docx
        glossary_path: path to terminology xlsx
        user_translations: dict of {cn_term: en_translation} for user-provided translations
        uncertain_terms: list of (cn_term, context) for terms that need user confirmation

    Returns:
        dict with keys: 'unmatched' (list of unmatched CN terms), 'uncertain' (terms needing confirmation)
    """
    if user_translations is None:
        user_translations = {}
    if uncertain_terms is None:
        uncertain_terms = []

    glossary = load_glossary(glossary_path)
    for cn, en in user_translations.items():
        glossary[cn] = en

    matcher = build_matcher(glossary)

    doc = Document(src_path)

    all_unmatched = []
    all_uncertain = []

    # Process paragraphs
    for i, para in enumerate(doc.paragraphs):
        if not para.text.strip():
            continue

        # Per-run replacement preserves individual run formatting
        replace_runs_in_place(para, matcher)

        # Collect unmatched terms from full paragraph text for reporting
        _, unmatched = translate_text(para.text, matcher)
        for term in unmatched:
            if len(term) >= 2 and term not in {'的', '了', '在', '和', '是', '与', '为',
                                                '以', '于', '被', '有', '这', '那', '个', '等'}:
                if term not in [u for u, _ in uncertain_terms]:
                    all_uncertain.append((term, para.text[:100]))
        if unmatched:
            all_unmatched.extend(unmatched)

        set_spacing(para, True)

    # Process tables
    for ti, tbl in enumerate(doc.tables):
        for ri, row in enumerate(tbl.rows):
            for ci, cell in enumerate(row.cells):
                orig_text = cell.text.strip()
                if not orig_text:
                    continue
                for pi, para in enumerate(cell.paragraphs):
                    if pi == 0:
                        replace_runs_in_place(para, matcher)
                        set_spacing(para, False)
                    else:
                        p_elem = para._p
                        p_elem.getparent().remove(p_elem)

    # Replace 【】bracket labels
    fix_cn_labels(doc)

    # Process text boxes — per-run replacement at XML level
    for txbx in doc.element.findall(f'.//{qn("w:txbxContent")}'):
        for p_elem in txbx.findall(qn('w:p')):
            text_runs_xml = []
            for r in p_elem.findall(qn('w:r')):
                t_els = r.findall(qn('w:t'))
                if any(t.text and t.text.strip() for t in t_els if t.text):
                    text_runs_xml.append(r)

            for r in text_runs_xml:
                t_els = r.findall(qn('w:t'))
                orig_text = ''.join(t.text or '' for t in t_els)
                trans_text, _ = translate_text(orig_text, matcher)
                if trans_text == orig_text:
                    continue

                for t in t_els:
                    r.remove(t)
                new_t = etree.SubElement(r, qn('w:t'))
                new_t.set(qn('xml:space'), 'preserve')
                new_t.text = trans_text

                rPr = r.find(qn('w:rPr'))
                if rPr is None:
                    rPr = etree.SubElement(r, qn('w:rPr'))
                _strip_shading(rPr)
                _force_tnr_12pt(rPr)

            # Set single spacing
            pPr = p_elem.find(qn('w:pPr'))
            if pPr is None:
                pPr = etree.SubElement(p_elem, qn('w:pPr'))
                p_elem.insert(0, pPr)
            sp = pPr.find(qn('w:spacing'))
            if sp is None:
                sp = etree.SubElement(pPr, qn('w:spacing'))
            sp.set(qn('w:line'), '240')
            sp.set(qn('w:lineRule'), 'auto')

    save_doc(doc, dst_path)

    return {
        'unmatched': list(dict.fromkeys(all_unmatched)),
        'uncertain': all_uncertain[:20],
    }


if __name__ == '__main__':
    # Quick test
    gpath = r'D:\WORK\客户信息\曹闵\曹闵术语表.xlsx'
    glossary = load_glossary(gpath)
    matcher = build_matcher(glossary)

    tests = [
        '曾供职于中海地产、万科地产、景瑞控股，拥有20余年房地产住宅产品研发管理工作经验',
        'CARR卡瑞装配式内装部品体系、Pincloud数字化云平台',
        '品宅·集栋装饰工程、品宅·一站美办',
        '核心合伙人向宠，联合创始人/首席执行官',
        '绿色建材选择与性能评估决策支持系统V1.0',
        '空间功能分区规划辅助设计系统 V1.0',
    ]

    for t in tests:
        out, unmatched = translate_text(t, matcher)
        print('IN :', t)
        print('OUT:', out)
        if unmatched:
            print('?? :', unmatched)
        print()
